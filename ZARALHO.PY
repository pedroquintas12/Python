import xlrd
import openpyxl
import pandas as pd
import re
from docx import Document
import mysql.connector
import os
from datetime import datetime, timedelta
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

db_config = {
    'host': 'localhost',
    'user': 'root',
    'password': '123456',
    'database': 'clientesdistribuicao'
}

data_do_dia = datetime.now()
data_formatada = data_do_dia.strftime('%Y%m%d')
nomeDoArquivo = ("impressao (1).xls")
nomeDoArquivoDocx = (data_formatada + "-" + "distribuição.docx")
FechamentoMes = (data_formatada +"-"+"FechamentoDoMes.docx" )


caminho_arquivo = r"C:\Users\pedro\OneDrive - LIG CONTATO DIÁRIO FORENSE\DISTRIBUIÇÃO\\" + nomeDoArquivo

if not os.path.exists(caminho_arquivo):
    data_do_dia -= timedelta(days=1)  
    data_formatada = data_do_dia.strftime('%Y%m%d')
    nomeDoArquivo = (data_formatada + "-" + "impressao.xls")
    caminho_arquivo = r"C:\Users\pedro\OneDrive - LIG CONTATO DIÁRIO FORENSE\DISTRIBUIÇÃO\DISTRIBUIÇÕES\\" + nomeDoArquivo
    nomeDoArquivoDocx = (data_formatada + "-" + "distribuição.docx")


print("Caminho do arquivo:", caminho_arquivo)


try:
    workbook = xlrd.open_workbook(caminho_arquivo)
    sheet = workbook.sheet_by_index(0)  
    print("Planilha carregada com sucesso.")

    # Colunas para extrair 
    colunas_alvo = [0, 2, 4, 5, 7, 10, 12]  

    dados_das_linhas = []

    # Defina o nome da coluna que você deseja tratar
    nome_coluna_tratamento = 'coluna_5'

    for row in range(1, sheet.nrows):  # Começando da segunda linha
        linha = {}
        for col_idx, col in enumerate(colunas_alvo):
            valor_celula = sheet.cell_value(row, col)
            nome_coluna = f'coluna_{col + 1}'
            contadorTotal =+ 1
        
            if col == 4:  # Se a coluna for a coluna 5
            # Remove todos os caracteres especiais, exceto "|"
             valor_celula = re.sub(r'[^\w\s|@ \w\s-]', '', valor_celula)
            # Remove números e a substring "(CNPJ)"
             valor_celula = re.sub(r'\d+|CNPJ', '', valor_celula, flags=re.IGNORECASE)
             valor_celula = re.sub(r'\d+|LTDA', '', valor_celula, flags=re.IGNORECASE)
             valor_celula = re.sub(r'\d+|CPF', '', valor_celula, flags=re.IGNORECASE)
             valor_celula = re.sub(r'\d+|EM PERNAMBUCO', '', valor_celula, flags=re.IGNORECASE)
            

             valor_celula = valor_celula.strip()
        
            linha[nome_coluna] = valor_celula
        dados_das_linhas.append(linha)
    
    # Imprime o valor após o tratamento dos caracteres especiais
    valor_tratado = linha.get(nome_coluna_tratamento)  # Usando .get() para tratar chaves ausentes
    print(f"Resultado após tratamento ({nome_coluna_tratamento}): {valor_tratado}")

    db_connection = mysql.connector.connect(**db_config)
    db_cursor = db_connection.cursor()

    nomes_pesquisa = [linha['coluna_5'] for linha in dados_das_linhas]

    for nome_pesquisa in nomes_pesquisa:
     names_to_search = [nome_pesquisa]
     
    
    for name in names_to_search:
        query = f"SELECT `CodEscritorio` FROM clientesdistribuicao.termosdistribuicao where `variacoes` LIKE '{name.strip()}%'"
        db_cursor.execute(query)
        result = db_cursor.fetchone()
    
        # Consumir todos os resultados da consulta anterior (se houver)
        if result is not None:
            for _ in result:
                pass
        resultado_query = result
    
        if resultado_query:
            codigo_escritorio = resultado_query[0]
            print(f"Resultado da pesquisa para '{name.strip()}': Código de Escritório = {codigo_escritorio}")
            codigo_escritorioSTG = str(codigo_escritorio)
            break  # Se encontrado, não é necessário continuar procurando
        else:
            print(f"Nenhum resultado encontrado para '{name.strip()}'")

    names_to_search_Client = resultado_query

    for nameClient in names_to_search_Client:
       query = f"SELECT `Nomedoescritorio` FROM clientesdistribuicao.clientes where `CodigoVSAP` = '{nameClient}'"
       db_cursor.execute(query)
       resultClient = db_cursor.fetchone()
    element_to_process = resultClient[0]  # Acesse o elemento desejado da tupla
    resultClient_modified = re.sub(r'[^\w\s|]', '', element_to_process)

    #cria o doc word
    doc = Document() 

    TextCodClient = (f"Código Escritório:  {codigo_escritorioSTG}")

    paragraph1 = doc.add_paragraph(resultClient)
    run1 = paragraph1.runs[0]
    font1 = run1.font
    font1.color.rgb = RGBColor(0x66, 0x99, 0xFF)  # Cor azul
    font1.size = Pt(26) 
    
    paragraph2 = doc.add_paragraph(TextCodClient)
    run2 = paragraph2.runs[0]
    font2 = run2.font
    font2.color.rgb = RGBColor(0x66, 0x99, 0xFF)  # Cor azul
    font2.size = Pt(26)

    border_paragraph = doc.add_paragraph('')
    border_paragraph.add_run('_______________________________________________________________________________________________')
    border_paragraph.runs[0].bold = True

# Corpo do e-mail
    email_template = (
    "Cliente: {resultClient2} \n\n"
    "Dados coletados:\n\n"
    "{dados}\n\n"
    "Atenciosamente,\n"
    "Lig Contato\n\n"
    "Detalhes:\n"
    "Data: {data}\n"
)

# Preencher o corpo do e-mail com os dados capturados
    dados_formatados = ""
    contador = 0
    for linha in dados_das_linhas:
        contador += 1
        dados_formatados += f"Distribuição: {contador}\n\n\n"
        dados_formatados += f"Número Processo: {linha['coluna_3']}\n\n"
        dados_formatados += f"Data distribuição: {linha['coluna_1']}\n\n"
        dados_formatados += f"Orgão: {linha['coluna_8']}\n\n"
        dados_formatados += f"Classe Judicial: {linha['coluna_13']}\n\n"
        dados_formatados += f"Polo Ativo: {linha['coluna_6']}\n\n"
        dados_formatados += f"Polo Passivo: {linha['coluna_5']}\n\n"
        dados_formatados += f"Link: {linha['coluna_11']}\n\n"
    dados_formatados += f"Total de distribuições: {contador}\n\n\n"


    email_texto = email_template.format(data=data_do_dia.strftime('%d/%m/%y'),
                                    dados=dados_formatados, resultClient2 = resultClient_modified)

    doc.add_paragraph(email_texto)
    
    caminho_arquivo_docx = r"C:\Users\pedro\OneDrive - LIG CONTATO DIÁRIO FORENSE\DISTRIBUIÇÃO\ARQUIVOS WORD DISTRIBUIÇÕES\\" + FechamentoMes
    doc.save(caminho_arquivo_docx)

   
except FileNotFoundError:
    print("Arquivo Excel não encontrado.")
except NameError:
    print("NOME NÃO ENCONTRADO NO BANCO DE DADOS")