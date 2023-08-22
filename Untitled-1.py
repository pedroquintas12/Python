import xlrd
import openpyxl
import pandas as pd
import re
from docx import Document
import mysql.connector
from datetime import datetime, timedelta

db_config = {
    'host': 'localhost',
    'user': 'root',
    'password': '123456',
    'database': 'clientesdistribuicao'
}

data_do_dia_anterior = datetime.now() - timedelta(days=1)
data_formatada = data_do_dia_anterior.strftime('%Y%m%d')
nomeDoArquivo = (data_formatada + "-" + "impressao.xls")
nomeDoArquivoDocx = (data_formatada + "-" + "distribuição.docx")

caminho_arquivo = r"C:\Users\pedro\OneDrive - LIG CONTATO DIÁRIO FORENSE\DISTRIBUIÇÃO\DISTRIBUIÇÕES\\" + nomeDoArquivo

print("Caminho do arquivo:", caminho_arquivo)


try:
    workbook = xlrd.open_workbook(caminho_arquivo)
    sheet = workbook.sheet_by_index(0)  
    print("Planilha carregada com sucesso.")

    # Colunas para extrair 
    colunas_alvo = [0, 2, 4, 5, 7, 10, 12]  

    dados_das_linhas = []

    # Defina o nome da coluna que você deseja tratar
    nome_coluna_tratamento = 'coluna_5'

    for row in range(1, sheet.nrows):  # Começando da segunda linha
        linha = {}
    for col_idx, col in enumerate(colunas_alvo):
        valor_celula = sheet.cell_value(row, col)
        nome_coluna = f'coluna_{col + 1}'
        
        if col == 4:  # Se a coluna for a coluna 5
            # Remove todos os caracteres especiais, exceto "|"
            valor_celula = re.sub(r'[^\w\s|]', '', valor_celula)
            # Remove números e a substring "(CNPJ)"
            valor_celula = re.sub(r'\d+|CNPJ', '', valor_celula, flags=re.IGNORECASE)
            valor_celula = valor_celula.strip()
        
        linha[nome_coluna] = valor_celula
    dados_das_linhas.append(linha)
    
    # Imprime o valor após o tratamento dos caracteres especiais
    valor_tratado = linha.get(nome_coluna_tratamento)  # Usando .get() para tratar chaves ausentes
    print(f"Resultado após tratamento ({nome_coluna_tratamento}): {valor_tratado}")

    db_connection = mysql.connector.connect(**db_config)
    db_cursor = db_connection.cursor()

    nomes_pesquisa = [linha['coluna_5'] for linha in dados_das_linhas]

    for nome_pesquisa in nomes_pesquisa:
     names_to_search = [nome_pesquisa]
     
    if '|' in nome_pesquisa:
        names_to_search = nome_pesquisa.split('|')
        print(f"Nenhum resultado encontrado para '{names_to_search}'")
    
    for name in names_to_search:
        query = f"SELECT `Cod Escritorio` FROM clientesdistribuicao.termosdistribuicao WHERE `Cod Escritorio` = '{name.strip()}'"
        db_cursor.execute(query)
        result = db_cursor.fetchone()
    
        # Consumir todos os resultados da consulta anterior (se houver)
        for _ in db_cursor.fetchall():
            pass
    
        resultado_query = db_cursor.fetchone()
    
        if resultado_query:
            codigo_escritorio = resultado_query[0]
            print(f"Resultado da pesquisa para '{name.strip()}': Código de Escritório = {codigo_escritorio}")
            break  # Se encontrado, não é necessário continuar procurando
    else:
        print(f"Nenhum resultado encontrado para '{nome_pesquisa}'")

    db_cursor.close()  # Feche o cursor após cada consulta

  #cria o doc word
    doc = Document()

# Adiciona informações ao documento
    doc.add_heading( 'distribuição', level=0,)

# Corpo do e-mail
    email_template = (
    "Prezado(a),\n\n"
    "Segue abaixo o relatório de distribuição referente ao dia anterior:\n\n"
    "Detalhes:\n"
    "Data: {data}\n"
    "Dados coletados:\n\n"
    "{dados}\n\n"
    "Atenciosamente,\n"
    "Lig Contato"
)

# Preencher o corpo do e-mail com os dados capturados
    dados_formatados = ""
    contador = 0
    for linha in dados_das_linhas:
        contador += 1
        dados_formatados += f"Distribuição: {contador}\n\n\n"
        dados_formatados += f"Numero Processo: {linha['coluna_3']}\n\n"
        dados_formatados += f"Data distribuição: {linha['coluna_1']}\n\n"
        dados_formatados += f"Orgão: {linha['coluna_8']}\n\n"
        dados_formatados += f"Classe Judicial: {linha['coluna_13']}\n\n"
        dados_formatados += f"Polo Ativo: {linha['coluna_6']}\n\n"
        dados_formatados += f"Polo Passivo: {linha['coluna_5']}\n\n"
        dados_formatados += f"Link: {linha['coluna_11']}\n\n"
       
    

    email_texto = email_template.format(data=data_do_dia_anterior.strftime('%d/%m/%y'),
                                    dados=dados_formatados)

    doc.add_paragraph(email_texto)

    caminho_arquivo_docx = r"C:\Users\pedro\OneDrive - LIG CONTATO DIÁRIO FORENSE\DISTRIBUIÇÃO\ARQUIVOS WORD DISTRIBUIÇÕES\\" + nomeDoArquivoDocx
    doc.save(caminho_arquivo_docx)

   
except FileNotFoundError:
    print("Arquivo Excel não encontrado.")
