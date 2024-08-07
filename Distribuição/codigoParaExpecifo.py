import re
from docx import Document
import mysql.connector
import os
from datetime import datetime, timedelta
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

db_config = {
    'host': 'localhost',
    'user': 'root',
    'password': '123456',
    'database': 'apidistribuicao'
}

data_do_dia = datetime.now()
data_formatada = data_do_dia.strftime('%Y-%m-%d')
nomeDoArquivo = (data_formatada + "-" + "impressao.xls")
nomeDoArquivoDocx = (data_formatada + "-" + "distribuição.docx")
try:
    db_connection = mysql.connector.connect(**db_config)
    db_cursor = db_connection.cursor()
    

    query = (
    "SELECT c.Cliente_VSAP as clienteVSAP, p.CodEscritorio, p.numeroProcesso, "
    "MAX(p.dataDistribuicao) as dataDistribuicao, "
    "p.orgaoJulgador, p.tipoDoProcesso, p.status, "
    "GROUP_CONCAT(DISTINCT a.nome ORDER BY a.nome SEPARATOR ', ') AS nomesAutores, "
    "GROUP_CONCAT(DISTINCT r.nome ORDER BY r.nome SEPARATOR ', ') AS nomesReus, "
    "GROUP_CONCAT(distinct l.linkDocumento  order by l.linkDocumento separator', ') as Link, "
    "p.uf, p.siglaSistema, MAX(p.instancia), p.tribunal "
    "FROM apidistribuicao.processo AS p "
    "LEFT JOIN apidistribuicao.clientes AS c ON p.CodEscritorio = c.CodEscritorio "
    "LEFT JOIN apidistribuicao.processo_autor AS a ON p.ID_processo = a.ID_processo "
    "LEFT JOIN apidistribuicao.processo_reu AS r ON p.ID_processo = r.ID_processo "
    "LEFT JOIN apidistribuicao.processo_docinicial as l ON p.ID_processo = l.ID_processo "
    f"WHERE DATE(p.data_insercao) = '{data_formatada}' "
    "and p.dataDistribuicao >= '2024-04-01'"
    "AND l.docPeticaoInicial = 0 "
    "AND p.CodEscritorio = 1378 "
    "AND p.status = 'S' "
    "AND p.uf IN ('PE', 'AL','CE','PI','BA','MA','PB','SE','RN') "
    "GROUP BY p.numeroProcesso, clienteVSAP, p.CodEscritorio, p.orgaoJulgador, p.tipoDoProcesso, p.uf, p.siglaSistema,p.tribunal;"
)


        
    db_cursor.execute(query)
    results =db_cursor.fetchall()
   # Inicialize a variável para rastrear o CodEscritorio atual
    cod_escritorio_atual = None
    doc = None  # Adicione esta linha para definir doc como None inicialmente
    contador_global = 0  # Inicialize o contador global

    # Dicionário para armazenar nomes de autores e réus por número de processo
    processos_dict = {}

    for idx, result in enumerate(results):
        cod_escritorio = result[1]  # Índice 1 corresponde à coluna CodEscritorio
        num_processo = result[2]  # Índice 2 corresponde ao número do processo

        if cod_escritorio != cod_escritorio_atual:
            # Se for um novo CodEscritorio, crie um novo arquivo Word
            if doc is not None:
                # Salve e feche o arquivo Word anterior, se existir
                doc.save(caminho_arquivo_docx)
                doc = None

                # Incremento do contador global
                contador_global += contador_local
                contador_local = 0  # Reinicialize o contador local para o próximo CodEscritorio

            # Atualize a variável para o CodEscritorio atual
            cod_escritorio_atual = cod_escritorio

            doc = Document()
            element_to_process = result[0] if result else ''
            resultClient_modified = re.sub(r'[^\w\s| ]', '', element_to_process)
            resultClient_modified = resultClient_modified.strip()

            # Crie um novo arquivo Word para o CodEscritorio atual
            caminho_arquivo_docx = os.path.join(
                "C:\\Users\\pedro\\OneDrive - LIG CONTATO DIÁRIO FORENSE\\DISTRIBUIÇÃO\\ARQUIVOS WORD DISTRIBUIÇÕES\\",
                f"{resultClient_modified}-{nomeDoArquivoDocx}"
            )
            TextCodClient = f"Código Escritório:  {result[1]}"

            paragraph1 = doc.add_paragraph(result[0])
            run1 = paragraph1.runs[0]
            font1 = run1.font
            font1.color.rgb = RGBColor(0x66, 0x99, 0xFF)  # Cor azul
            font1.size = Pt(26)

            paragraph2 = doc.add_paragraph(TextCodClient)
            run2 = paragraph2.runs[0]
            font2 = run2.font
            font2.color.rgb = RGBColor(0x66, 0x99, 0xFF)  # Cor azul
            font2.size = Pt(26)

            border_paragraph = doc.add_paragraph('')
            border_paragraph.add_run('_______________________________________________________________________________________________')
            border_paragraph.runs[0].bold = True

            # Inicialize o contador local para cada CodEscritorio
            contador_local = 0

        # Verifique se o número do processo já está no dicionário
        if num_processo in processos_dict:
            # Adicione o nomeAutor e nomeReu ao dicionário existente
            processos_dict[num_processo]['nomeAutor'].append(result[7])
            processos_dict[num_processo]['nomeReu'].append(result[8])
        else:
            # Crie um novo item no dicionário para o número do processo
            processos_dict[num_processo] = {'nomeAutor': [result[7]], 'nomeReu': [result[8]]}

        # Se este não for o último resultado e o próximo tiver o mesmo número de processo,
        # continue acumulando nomes no dicionário até encontrar um número de processo diferente
        if idx != len(results) - 1 and results[idx + 1][2] == num_processo:
            continue

        query2 = ("SELECT c.Cliente_VSAP as clienteVSAP, p.CodEscritorio, p.numeroProcesso, p.status, MAX(p.dataDistribuicao) as dataDistribuicao, " 
                    "MAX(p.orgaoJulgador) as orgaoJulgador, MAX(p.tipoDoProcesso) as tipoDoProcesso, "
                    "MAX(a.nome) AS nomeAutor, MAX(r.nome) AS nomeReu, MAX(l.link) as link, "
                    "MAX(date(p.data_insercao)) as Data "
                    "FROM apidistribuicao.processo AS p "
                    "LEFT JOIN apidistribuicao.clientes AS c ON p.CodEscritorio = c.CodEscritorio "
                    "LEFT JOIN apidistribuicao.processo_autor AS a ON p.ID_processo = a.ID_processo "
                    "LEFT JOIN apidistribuicao.processo_reu AS r ON p.ID_processo = r.ID_processo "
                    "LEFT JOIN apidistribuicao.processo_link AS l ON p.ID_processo = l.ID_processo "
                    "WHERE p.deleted = 0 "
                        "and p.dataDistribuicao >= '2024-04-01'"
                        f"AND DATE (p.data_insercao) >= '{data_formatada}' "
                        f"AND c.CodEscritorio = {cod_escritorio_atual} "
                        "AND p.status = 'S' "
                        "AND p.uf IN ('PE', 'AL','CE','PI','BA','MA','PB','SE','RN') "
                    "GROUP BY p.numeroProcesso, clienteVSAP, p.CodEscritorio;") 
        
        db_cursor.execute(query2)
        results_rows =db_cursor.fetchall()
        
        # Corpo do e-mail
        email_template = (
            "Cliente: {resultClient2} \n\n"
            "Dados coletados:\n\n"
            "{dados}\n\n"
        )

        dados_formatados = ""
        contador_local += 1
        
        # Preencher o corpo do e-mail com os dados capturados
        dados_formatados += f"Distribuição: {contador_local} de {len(results_rows)} \n\n\n"
        dados_formatados += f"Tribunal: {result[13]} \n\n"
        dados_formatados += f"UF/Instância/Comarca:    {result[10]}/{result[12]}/{result[11]} \n\n"
        dados_formatados += f"Número Processo: {num_processo}\n\n"
        data_distribuicao = datetime.strptime(str(result[3]), '%Y-%m-%d').strftime('%d/%m/%Y')
        dados_formatados += f"Data distribuição: {data_distribuicao}\n\n"
        dados_formatados += f"Orgão: {result[4]}\n\n"
        dados_formatados += f"Classe Judicial: {result[5]}\n\n"
        if processos_dict[num_processo]['nomeAutor'] is not None:
            dados_formatados += f"Polo Ativo: {', '.join((filter(None, processos_dict [num_processo]['nomeAutor'])))}\n\n"
        else:
            dados_formatados += "Polo ativo: NULL"

        if processos_dict[num_processo]['nomeReu'] is not None:
            dados_formatados += f"Polo Passivo: {', '.join(filter(None, processos_dict[num_processo]['nomeReu']))}\n\n"
        else:
            dados_formatados += "Polo Passivo: [Nenhum dado disponível]\n\n"       
        
        dados_formatados += f"Link: {result[9]}\n\n"

        email_texto = email_template.format(data=data_do_dia.strftime('%d/%m/%y'),
                                            dados=dados_formatados, resultClient2=resultClient_modified)

        doc.add_paragraph(email_texto)

    # Salve o último arquivo Word
    if doc is not None:
        paragraph_atenciosamente = doc.add_paragraph("Atenciosamente,")
        paragraph_atenciosamente = doc.add_paragraph("Lig Contato\n")
        dataComBarra = data_do_dia.strftime('%d/%m/%Y')
        paragraph_atenciosamente = doc.add_paragraph(f"Data: {dataComBarra}")
        
        doc.save(caminho_arquivo_docx)
        contador_global += contador_local  # Incremento final do contador global

    # Exiba o contador global ao final do processamento
    print(f"Total de Distribuições: {contador_global}")

except NameError:
    print("NOME NÃO ENCONTRADO NO BANCO DE DADOS")