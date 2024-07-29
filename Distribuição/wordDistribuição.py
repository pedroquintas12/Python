import re
from docx import Document
import mysql.connector
import os
from datetime import datetime
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL

db_config = {
    'host': '26.191.28.12',
    'port': '3306',
    'user': 'pedro',
    'password': '123456',
    'database': 'apidistribuicao'
}

data_do_dia = datetime.now()
data_formatada = data_do_dia.strftime('%Y-%m-%d')
nomeDoArquivoDocx = f"{data_formatada}-distribuição.docx"

try:
    db_connection = mysql.connector.connect(**db_config)
    db_cursor = db_connection.cursor()

    query_cod_escritorio = (
        "SELECT cod_escritorio FROM processo "
        "WHERE DATE(data_insercao) = %s "
        "GROUP BY cod_escritorio"
    )
    db_cursor.execute(query_cod_escritorio, (data_formatada,))
    cod_escritorios = db_cursor.fetchall()
    cod_escritorios = [row[0] for row in cod_escritorios]

    if not cod_escritorios:
        print("Nenhum código de escritório encontrado para a data especificada.")
        exit()

    for cod_escritorio in cod_escritorios:
        query = (
            "SELECT c.Cliente_VSAP as clienteVSAP, p.Cod_escritorio, p.numero_processo, "
            "MAX(p.data_distribuicao) as data_distribuicao, "
            "p.orgao_julgador, p.tipo_processo, p.status, "
            "GROUP_CONCAT(DISTINCT a.nome ORDER BY a.nome SEPARATOR ', ') AS nomesAutores, "
            "GROUP_CONCAT(DISTINCT r.nome ORDER BY r.nome SEPARATOR ', ') AS nomesReus, "
            "GROUP_CONCAT(distinct l.link_documento order by l.link_documento separator ' | ') as Link, "
            "p.uf, p.sigla_sistema, MAX(p.instancia), p.tribunal "
            "FROM apidistribuicao.processo AS p "
            "LEFT JOIN apidistribuicao.clientes AS c ON p.Cod_escritorio = c.Cod_escritorio "
            "LEFT JOIN apidistribuicao.processo_autor AS a ON p.ID_processo = a.ID_processo "
            "LEFT JOIN apidistribuicao.processo_reu AS r ON p.ID_processo = r.ID_processo "
            "LEFT JOIN apidistribuicao.processo_docinicial as l ON p.ID_processo = l.ID_processo "
            "WHERE DATE(p.data_insercao) = %s "
            "AND l.doc_peticao_inicial = 0 "
            "AND p.Cod_escritorio = %s "
            "AND p.status = 'S' "
            "GROUP BY p.numero_processo, clienteVSAP, p.Cod_escritorio, p.orgao_julgador, p.tipo_processo, p.uf, p.sigla_sistema, p.tribunal;"
        )
        
        db_cursor.execute(query, (data_formatada, cod_escritorio))
        results = db_cursor.fetchall()

        cod_escritorio_atual = None
        doc = None  
        contador_global = 0  
   
        processos_dict = {}

        for idx, result in enumerate(results):
            cod_escritorio = result[1] 
            num_processo = result[2]  

            if cod_escritorio != cod_escritorio_atual:
                if doc is not None:
                    doc.save(caminho_arquivo_docx)
                    doc = None

                    contador_global += contador_local
                    contador_local = 0  

                cod_escritorio_atual = cod_escritorio

                doc = Document()
                element_to_process = result[0] if result else ''

                caminho_arquivo_docx = os.path.join(
                    "C:\\Users\\pedro\\OneDrive - LIG CONTATO DIÁRIO FORENSE\\DISTRIBUIÇÃO\\ARQUIVOS WORD DISTRIBUIÇÕES\\",
                    f"{element_to_process}-{nomeDoArquivoDocx}"
                )
                TextCodClient = f"Código Escritório:  {result[1]}"

                paragraph1 = doc.add_paragraph(result[0])
                run1 = paragraph1.runs[0]
                font1 = run1.font
                font1.color.rgb = RGBColor(0x66, 0x99, 0xFF)  # Cor azul
                font1.size = Pt(26)

                paragraph2 = doc.add_paragraph(TextCodClient)
                run2 = paragraph2.runs[0]
                font2 = run2.font
                font2.color.rgb = RGBColor(0x66, 0x99, 0xFF)  # Cor azul
                font2.size = Pt(26)

                border_paragraph = doc.add_paragraph('')
                border_paragraph.add_run('_______________________________________________________________________________________________')
                border_paragraph.runs[0].bold = True

                contador_local = 0

            if num_processo in processos_dict:
                processos_dict[num_processo]['nomeAutor'].append(result[7])
                processos_dict[num_processo]['nomeReu'].append(result[8])
            else:
                processos_dict[num_processo] = {'nomeAutor': [result[7]], 'nomeReu': [result[8]]}

            if idx != len(results) - 1 and results[idx + 1][2] == num_processo:
                continue

            query2 = (
                "SELECT c.Cliente_VSAP as clienteVSAP, p.Cod_escritorio, p.numero_processo, p.status, MAX(p.data_distribuicao) as data_distribuicao, " 
                "MAX(p.orgao_julgador) as orgao_julgador, MAX(p.tipo_processo) as tipo_processo, "
                "MAX(a.nome) AS nomeAutor, MAX(r.nome) AS nomeReu, MAX(l.link) as link, "
                "MAX(DATE(p.data_insercao)) as Data "
                "FROM apidistribuicao.processo AS p "
                "LEFT JOIN apidistribuicao.clientes AS c ON p.Cod_escritorio = c.Cod_escritorio "
                "LEFT JOIN apidistribuicao.processo_autor AS a ON p.ID_processo = a.ID_processo "
                "LEFT JOIN apidistribuicao.processo_reu AS r ON p.ID_processo = r.ID_processo "
                "LEFT JOIN apidistribuicao.processo_link AS l ON p.ID_processo = l.ID_processo "
                "WHERE p.deleted = 0 "
                "AND DATE(p.data_insercao) >= %s "
                "AND c.Cod_escritorio = %s "
                "AND p.status = 'S' "
                "GROUP BY p.numero_processo, clienteVSAP, p.Cod_escritorio;"
            )
            
            db_cursor.execute(query2, (data_formatada, cod_escritorio_atual))
            results_rows = db_cursor.fetchall()

            email_template = (
                "Cliente: {resultClient2} \n\n"
                "Dados coletados:\n\n"
                "{dados}\n\n"
            )

            dados_formatados = ""
            contador_local += 1
        
            dados_formatados += f"Distribuição: {contador_local} de {len(results_rows)} \n\n\n"
            dados_formatados += f"Tribunal: {result[13]} \n\n"
            dados_formatados += f"UF/Instância/Comarca:    {result[10]}/{result[12]}/{result[11]} \n\n"
            dados_formatados += f"Número Processo: {num_processo}\n\n"
            data_distribuicao = datetime.strptime(str(result[3]), '%Y-%m-%d').strftime('%d/%m/%Y')
            dados_formatados += f"Data distribuição: {data_distribuicao}\n\n"
            dados_formatados += f"Orgão: {result[4]}\n\n"
            dados_formatados += f"Classe Judicial: {result[5]}\n\n"
            if processos_dict[num_processo]['nomeAutor'] is not None:
                dados_formatados += f"Polo Ativo: {', '.join((filter(None, processos_dict [num_processo]['nomeAutor'])))}\n\n"
            else:
                dados_formatados += "Polo ativo: [Nenhum dado disponível]"

            if processos_dict[num_processo]['nomeReu'] is not None:
                dados_formatados += f"Polo Passivo: {', '.join(filter(None, processos_dict[num_processo]['nomeReu']))}\n\n"
            else:
                dados_formatados += "Polo Passivo: [Nenhum dado disponível]\n\n"       
            
            dados_formatados += f"Link: {result[9]}\n\n"

            email_texto = email_template.format(data=data_do_dia.strftime('%d/%m/%y'),
                                                dados=dados_formatados, resultClient2=element_to_process)

            doc.add_paragraph(email_texto)

        if doc is not None:
            paragraph_atenciosamente = doc.add_paragraph("Atenciosamente,")
            paragraph_atenciosamente = doc.add_paragraph("Lig Contato\n")
            dataComBarra = data_do_dia.strftime('%d/%m/%Y')
            paragraph_atenciosamente = doc.add_paragraph(f"Data: {dataComBarra}")
            
            doc.save(caminho_arquivo_docx)
            contador_global += contador_local  

        print(f"Total de Distribuições para {element_to_process}:\n {contador_global}")

except NameError:
    print("NOME NÃO ENCONTRADO NO BANCO DE DADOS")